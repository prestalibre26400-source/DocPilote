import hashlib
import os
import secrets
import sqlite3
import subprocess
import tempfile
import time
from contextlib import contextmanager
from datetime import datetime, timedelta, timezone

from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from pypdf import PdfReader
import requests

import docx as docx_lib
from odf import teletype
from odf.opendocument import load as load_odf
from odf.text import P as OdfParagraph
from striprtf.striprtf import rtf_to_text
import stripe

MISTRAL_API_KEY = os.environ.get("MISTRAL_API_KEY", "")
MISTRAL_MODEL = os.environ.get("MISTRAL_MODEL", "mistral-small-latest")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "")
IP_HASH_SALT = os.environ.get("IP_HASH_SALT", "")

# ---------------------------------------------------------------------------
# Modèle économique : freemium + abonnement mensuel via Stripe
# ---------------------------------------------------------------------------
STRIPE_SECRET_KEY = os.environ.get("STRIPE_SECRET_KEY", "")
STRIPE_WEBHOOK_SECRET = os.environ.get("STRIPE_WEBHOOK_SECRET", "")
STRIPE_PRICE_ID = os.environ.get("STRIPE_PRICE_ID", "")
stripe.api_key = STRIPE_SECRET_KEY
FREE_MONTHLY_QUOTA = int(os.environ.get("DOCPILOTE_FREE_QUOTA", "20"))
MONTHLY_PRICE_EUR = int(os.environ.get("DOCPILOTE_MONTHLY_PRICE_EUR", "12"))
CHECKOUT_SUCCESS_URL = "https://docpilote.prestalibre.org/licence-activee.html?session_id={CHECKOUT_SESSION_ID}"
CHECKOUT_CANCEL_URL = "https://docpilote.prestalibre.org/#pricing"

app = FastAPI(title="DocPilote API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

MAX_FILE_BYTES = 20 * 1024 * 1024
MAX_TEXT_CHARS = 60000

# ---------------------------------------------------------------------------
# Version courante distribuée + emplacement des .deb téléchargeables
# ---------------------------------------------------------------------------

CURRENT_VERSION = "0.13.7"
CURRENT_VERSION_WINDOWS = "1.0.6"
DOWNLOADS_DIR = "/var/www/docpilote/downloads"

# ---------------------------------------------------------------------------
# Base de données de suivi (téléchargements + usage) — RGPD : aucune IP en
# clair n'est jamais stockée, seulement un hash salé non réversible.
# ---------------------------------------------------------------------------

DB_PATH = os.environ.get("DOCPILOTE_DB_PATH", "/var/www/docpilote-api/docpilote.db")


def _hash_ip(ip: str) -> str:
    if not ip:
        return ""
    return hashlib.sha256(f"{IP_HASH_SALT}:{ip}".encode("utf-8")).hexdigest()[:16]


@contextmanager
def db_conn():
    conn = sqlite3.connect(DB_PATH, timeout=10)
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def init_db():
    with db_conn() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS downloads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                version TEXT NOT NULL,
                ip_hash TEXT,
                user_agent TEXT,
                country TEXT,
                created_at TEXT NOT NULL
            )
            """
        )
        # Colonne ajoutee apres la creation initiale de la table (deploiements
        # existants) : ALTER TABLE ignore silencieusement si elle existe deja.
        try:
            conn.execute("ALTER TABLE downloads ADD COLUMN country TEXT")
        except sqlite3.OperationalError:
            pass
        # Geoloc ville (Cloudflare "Add visitor location headers") : ajoutee
        # apres coup, migration best-effort colonne par colonne.
        for ddl in (
            "ALTER TABLE downloads ADD COLUMN city TEXT",
            "ALTER TABLE downloads ADD COLUMN region TEXT",
            "ALTER TABLE downloads ADD COLUMN latitude REAL",
            "ALTER TABLE downloads ADD COLUMN longitude REAL",
        ):
            try:
                conn.execute(ddl)
            except sqlite3.OperationalError:
                pass
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS usage_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                endpoint TEXT NOT NULL,
                plugin TEXT,
                ext TEXT,
                success INTEGER NOT NULL,
                error_msg TEXT,
                duration_ms INTEGER,
                ip_hash TEXT,
                created_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_downloads_created ON downloads(created_at)"
        )
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_usage_created ON usage_events(created_at)"
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS licenses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                license_key TEXT NOT NULL UNIQUE,
                email TEXT,
                status TEXT NOT NULL DEFAULT 'active',
                stripe_customer_id TEXT,
                stripe_subscription_id TEXT,
                current_period_end TEXT,
                created_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS license_calls (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                license_key TEXT,
                created_at TEXT NOT NULL
            )
            """
        )
        conn.execute(
            "CREATE INDEX IF NOT EXISTS idx_license_calls_key ON license_calls(license_key, created_at)"
        )
        # Migration best-effort : colonnes de liaison machine (ajoutees apres
        # coup, sans casser une base existante). SQLite ne supporte pas
        # "ADD COLUMN IF NOT EXISTS" nativement -> try/except par colonne.
        for ddl in (
            "ALTER TABLE licenses ADD COLUMN device_id TEXT",
            "ALTER TABLE licenses ADD COLUMN device_bound_at TEXT",
            "ALTER TABLE licenses ADD COLUMN device_changed_at TEXT",
        ):
            try:
                conn.execute(ddl)
            except sqlite3.OperationalError:
                pass  # colonne deja presente


init_db()


def log_download(
    version: str,
    ip: str,
    user_agent: str,
    country: str = "",
    city: str = "",
    region: str = "",
    latitude: str = "",
    longitude: str = "",
) -> None:
    def _to_float(v: str):
        try:
            return float(v) if v else None
        except ValueError:
            return None

    with db_conn() as conn:
        conn.execute(
            """
            INSERT INTO downloads (version, ip_hash, user_agent, country, city, region, latitude, longitude, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                version,
                _hash_ip(ip),
                (user_agent or "")[:200],
                (country or "")[:2].upper() or None,
                (city or "")[:100] or None,
                (region or "")[:100] or None,
                _to_float(latitude),
                _to_float(longitude),
                datetime.now(timezone.utc).isoformat(),
            ),
        )


def log_usage(endpoint: str, plugin: str, ext: str, success: bool, error_msg: str, duration_ms: int, ip: str) -> None:
    try:
        with db_conn() as conn:
            conn.execute(
                """
                INSERT INTO usage_events (endpoint, plugin, ext, success, error_msg, duration_ms, ip_hash, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    endpoint,
                    plugin,
                    ext,
                    1 if success else 0,
                    (error_msg or "")[:300],
                    duration_ms,
                    _hash_ip(ip),
                    datetime.now(timezone.utc).isoformat(),
                ),
            )
    except Exception:  # noqa: BLE001
        pass  # le tracking ne doit jamais casser une requête réelle


def _client_ip(request: Request) -> str:
    fwd = request.headers.get("x-forwarded-for", "")
    if fwd:
        return fwd.split(",")[0].strip()
    return request.client.host if request.client else ""


# ---------------------------------------------------------------------------
# Licences (freemium + abonnement mensuel Stripe)
#
# Modele : chaque poste peut appeler l'API sans cle (mode gratuit, plafonne
# a FREE_MONTHLY_QUOTA appels "gratuits" par mois, compte par hash d'IP,
# best-effort car une IP peut changer). Une licence payante (cle generee a
# l'achat Stripe) leve le plafond : usage illimite tant que l'abonnement est
# actif. La cle ne transite jamais en clair dans les logs; seule sa presence
# et sa validite sont verifiees ici.
# ---------------------------------------------------------------------------


def generate_license_key() -> str:
    raw = secrets.token_hex(10).upper()
    return f"DP-{raw[0:4]}-{raw[4:8]}-{raw[8:12]}-{raw[12:16]}"


def get_license(license_key: str):
    if not license_key:
        return None
    with db_conn() as conn:
        row = conn.execute(
            "SELECT license_key, email, status, current_period_end FROM licenses WHERE license_key = ?",
            (license_key,),
        ).fetchone()
    return row


def license_is_valid(license_key: str) -> bool:
    row = get_license(license_key)
    if not row:
        return False
    _, _, status, current_period_end = row
    if status != "active":
        return False
    if current_period_end:
        try:
            end = datetime.fromisoformat(current_period_end)
            if end < datetime.now(timezone.utc):
                return False
        except ValueError:
            pass
    return True


# ---------------------------------------------------------------------------
# Liaison licence <-> machine
#
# Le client envoie un device_id (hash local, jamais un identifiant materiel
# en clair) a l'activation et a chaque appel. Premiere utilisation = liaison
# automatique. Si un device_id different apparait pour la meme licence :
# refus, sauf si aucun changement recent n'a eu lieu (1 changement autorise
# toutes les DEVICE_CHANGE_COOLDOWN_DAYS jours, pour ne pas coincer un
# client qui change de PC legitimement). Objectif : eviter qu'une seule cle
# a 12E/mois soit partagee en illimite sur un nombre quelconque de machines.
# ---------------------------------------------------------------------------

DEVICE_CHANGE_COOLDOWN_DAYS = 30


def check_device_binding(license_key: str, device_id: str) -> None:
    """Leve HTTPException 409 si device_id ne correspond pas a l'appareil
    deja lie a cette licence et que le delai de changement n'est pas ecoule.
    Ne fait rien si device_id est vide (client ancien, sans support device
    -> pas d'application retroactive, evite de casser les licences deja
    activees avant ce fix)."""
    if not device_id:
        return
    with db_conn() as conn:
        row = conn.execute(
            "SELECT device_id, device_bound_at, device_changed_at FROM licenses WHERE license_key = ?",
            (license_key,),
        ).fetchone()
    if not row:
        return
    bound_device, bound_at, changed_at = row
    now = datetime.now(timezone.utc)

    if not bound_device:
        with db_conn() as conn:
            conn.execute(
                "UPDATE licenses SET device_id = ?, device_bound_at = ? WHERE license_key = ?",
                (device_id, now.isoformat(), license_key),
            )
        return

    if bound_device == device_id:
        return

    # Point de reference pour le cooldown : le dernier changement s'il y en
    # a eu un, sinon la liaison initiale (sinon le premier changement ne
    # serait jamais bloque, faute de device_changed_at existant).
    reference = changed_at or bound_at
    if reference:
        try:
            last_change = datetime.fromisoformat(reference)
            if last_change.tzinfo is None:
                last_change = last_change.replace(tzinfo=timezone.utc)
            if (now - last_change).days < DEVICE_CHANGE_COOLDOWN_DAYS:
                raise HTTPException(
                    409,
                    "Cette licence est deja active sur un autre appareil. "
                    "Le changement d'appareil est limite a 1 fois par mois — "
                    "contactez le support si besoin.",
                )
        except ValueError:
            pass

    with db_conn() as conn:
        conn.execute(
            "UPDATE licenses SET device_id = ?, device_changed_at = ? WHERE license_key = ?",
            (device_id, now.isoformat(), license_key),
        )


def free_calls_this_month(ip_hash: str) -> int:
    month_start = datetime.now(timezone.utc).replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    with db_conn() as conn:
        row = conn.execute(
            """
            SELECT COUNT(*) FROM usage_events
            WHERE ip_hash = ? AND success = 1 AND created_at >= ?
            """,
            (ip_hash, month_start.isoformat()),
        ).fetchone()
    return row[0] if row else 0


def log_license_call(license_key: str) -> None:
    if not license_key:
        return
    try:
        with db_conn() as conn:
            conn.execute(
                "INSERT INTO license_calls (license_key, created_at) VALUES (?, ?)",
                (license_key, datetime.now(timezone.utc).isoformat()),
            )
    except Exception:  # noqa: BLE001
        pass


def check_quota(request: Request, license_key: str, device_id: str = "") -> None:
    """Leve HTTP 402 si le quota gratuit mensuel est depasse et qu'aucune
    licence valide n'est fournie. Une licence valide = usage illimite, sous
    reserve de la liaison machine (leve HTTP 409 via DeviceConflictError,
    convertie en HTTPException par les endpoints appelants)."""
    if license_key and license_is_valid(license_key):
        check_device_binding(license_key, device_id)
        log_license_call(license_key)
        return

    ip_hash = _hash_ip(_client_ip(request))
    used = free_calls_this_month(ip_hash)
    if used >= FREE_MONTHLY_QUOTA:
        raise HTTPException(
            402,
            f"Quota gratuit mensuel atteint ({FREE_MONTHLY_QUOTA} analyses/mois). "
            f"Passez a la version illimitee : https://docpilote.prestalibre.org/#pricing",
        )


# ---------------------------------------------------------------------------
# Auth admin (Basic Auth)
# ---------------------------------------------------------------------------

security = HTTPBasic()


def require_admin(credentials: HTTPBasicCredentials = Depends(security)):
    valid_user = secrets.compare_digest(credentials.username, "admin")
    valid_pass = secrets.compare_digest(credentials.password, ADMIN_PASSWORD or "")
    if not (valid_user and valid_pass and ADMIN_PASSWORD):
        raise HTTPException(
            status_code=401,
            detail="Authentification requise",
            headers={"WWW-Authenticate": "Basic"},
        )
    return True


# ---------------------------------------------------------------------------
# Prompts pour les plugins "un fichier -> un texte"
#
# Multilingue : les instructions donnees au modele restent en francais (seule
# langue maintenue dans le code), mais la LANGUE DE SORTIE demandee au modele
# est parametrable via `lang` ("fr" par defaut, retro-compatible avec les
# clients existants qui n'envoient pas ce champ). Un LLM suit tres bien une
# consigne "redige ta reponse en anglais" meme si l'instruction elle-meme est
# en francais -> pas besoin de dupliquer chaque prompt par langue.
# ---------------------------------------------------------------------------

LANG_LABELS = {
    "fr": "français",
    "en": "anglais",
}


def lang_label(lang: str) -> str:
    return LANG_LABELS.get(lang, LANG_LABELS["fr"])


PLUGIN_PROMPTS = {
    "summarize": (
        "Résume ce document en {label}, de façon claire et structurée "
        "(points clés en puces si pertinent). Reste concis (200-300 mots max).\n\n"
        "Document:\n{text}"
    ),
    "respond": (
        "Tu es un assistant qui aide un professionnel à répondre à un document reçu "
        "(réclamation, courrier, demande...). Analyse ce document et rédige une "
        "proposition de réponse professionnelle et adaptée en {label}.\n"
        "- Si c'est une réclamation : propose une réponse empathique et posée.\n"
        "- Si c'est une demande d'information : réponds aux points soulevés.\n"
        "- Si un refus argumenté est plus approprié, rédige-le clairement, sans être cassant.\n"
        "Termine par une formule de politesse professionnelle adaptée à la langue de sortie. "
        "Indique en une ligne, avant le texte de réponse, le type de réponse choisi (mail / "
        "lettre / refus argumenté / demande d'information), dans cette même langue de sortie.\n\n"
        "Document reçu:\n{text}"
    ),
    "explain": (
        "Explique ce document en {label} de façon claire et structurée, pas comme un "
        "résumé générique mais comme une vraie explication utile pour quelqu'un qui doit "
        "agir dessus. Structure ta réponse en sections (titres de sections traduits dans la "
        "langue de sortie demandée) :\n"
        "**Résumé** (2-3 phrases)\n"
        "**Obligations** (ce que chaque partie doit faire)\n"
        "**Risques** (ce qui pourrait mal se passer)\n"
        "**Dates et délais importants**\n"
        "**Clauses importantes à retenir**\n"
        "Si une section est vide ou non pertinente, dis-le brièvement plutôt que de l'omettre.\n\n"
        "Document:\n{text}"
    ),
    "decide": (
        "Analyse ce document et réponds uniquement à la question, en {label} : "
        "quelles actions concrètes la personne qui reçoit ce document doit-elle entreprendre, "
        "et dans quel délai ? Liste les actions sous forme de checklist avec des ✓, "
        "en précisant les délais et échéances quand ils sont mentionnés. "
        "Sois concret et actionnable, pas descriptif. Si aucune action n'est requise, dis-le clairement.\n\n"
        "Document:\n{text}"
    ),
    "risks": (
        "Analyse ce document comme le ferait un auditeur des risques, en {label}. Identifie les "
        "points problématiques ou déséquilibrés pour la partie qui le reçoit : pénalités élevées, "
        "clauses de renouvellement automatique, absence de limitation de responsabilité, "
        "délais trop courts, obligations disproportionnées, zones d'ambiguïté, etc. "
        "Présente chaque risque identifié avec ⚠️ suivi d'une courte explication. "
        "Si le document ne présente pas de risque notable, dis-le clairement plutôt que d'en inventer.\n\n"
        "Document:\n{text}"
    ),
    "checklist": (
        "Transforme ce document en checklist pratique et actionnable en {label}. "
        "Identifie le type de document (cahier des charges, notice, procédure...) et "
        "génère une liste à cocher (☐) des éléments à vérifier, faire ou respecter, "
        "organisée par catégories si pertinent. Reste concret, pas de généralités.\n\n"
        "Document:\n{text}"
    ),
}

EXTRACT_PROMPT = (
    "Extrait les données structurées importantes de ce document (facture, contrat, "
    "formulaire...) et retourne-les STRICTEMENT au format JSON valide, sans texte autour, "
    "sans balises markdown. Utilise des clés en {label}, snake_case. Inclue par exemple "
    "(si présents) : émetteur, destinataire, date, montants, échéances, références, "
    "articles, obligations. N'invente aucune donnée absente du document.\n\n"
    "Document:\n{text}"
)

COMPARE_PROMPT = (
    "Compare ces deux versions d'un même document, en {label}. Identifie les différences "
    "significatives (articles modifiés, montants changés, clauses ajoutées ou supprimées, "
    "durées modifiées...). Présente le résultat de façon structurée, par exemple :\n"
    "**Article X modifié**\n"
    "Avant : ...\n"
    "Après : ...\n"
    "Ne liste pas les différences de mise en forme insignifiantes. Si les documents sont "
    "quasi identiques, dis-le clairement.\n\n"
    "--- Document 1 ---\n{text1}\n\n--- Document 2 ---\n{text2}"
)

ASK_PROMPT = (
    "Tu réponds à une question précise sur le document ci-dessous, en {label}. "
    "Réponds uniquement à partir du contenu du document. Si la réponse ne s'y trouve pas, "
    "dis-le clairement plutôt que d'inventer.\n\n"
    "Question : {question}\n\n"
    "Document:\n{text}"
)

TRANSFORM_TARGETS = {
    "mail": "un e-mail professionnel reprenant le contenu et l'intention du document",
    "procedure": "une procédure étape par étape, claire et actionnable",
    "synthese": "une synthèse exécutive courte (10 lignes maximum) pour une direction",
    "tableau": (
        "un tableau au format Markdown (colonnes séparées par |) reprenant les "
        "informations clés du document de façon structurée"
    ),
    "presentation": "un plan de présentation orale (slides) avec titres et points clés par slide",
    "compte_rendu": "un compte rendu formel structuré (contexte, points abordés, décisions, actions)",
}


SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".odt", ".txt", ".md", ".rtf",
    ".doc", ".xls", ".ppt", ".xlsx", ".pptx", ".ods", ".odp",
}

# Formats anciens/binaires convertis via LibreOffice headless -> PDF,
# puis extraits avec le même chemin que les PDF natifs.
LIBREOFFICE_EXTENSIONS = {".doc", ".xls", ".ppt", ".xlsx", ".pptx", ".ods", ".odp"}
LIBREOFFICE_TIMEOUT = 45


def _convert_with_libreoffice(content: bytes, ext: str) -> bytes:
    """Convertit un fichier Office (.doc/.xls/.ppt/.xlsx/.pptx/.ods/.odp) en
    PDF via LibreOffice headless, avec un profil utilisateur isolé et
    jetable pour ne jamais entrer en conflit entre requêtes concurrentes."""
    with tempfile.TemporaryDirectory() as workdir, tempfile.TemporaryDirectory() as profile_dir:
        src_path = os.path.join(workdir, f"input{ext}")
        with open(src_path, "wb") as f:
            f.write(content)

        env = os.environ.copy()
        env["HOME"] = workdir

        try:
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--norestore",
                    f"-env:UserInstallation=file://{profile_dir}",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    workdir,
                    src_path,
                ],
                capture_output=True,
                timeout=LIBREOFFICE_TIMEOUT,
                text=True,
                cwd=workdir,
                env=env,
            )
        except subprocess.TimeoutExpired:
            raise HTTPException(
                504, "La conversion du document a pris trop de temps (fichier trop complexe ?)"
            )

        pdf_path = os.path.join(workdir, "input.pdf")
        if result.returncode != 0 or not os.path.isfile(pdf_path):
            raise HTTPException(
                400,
                f"Conversion impossible pour ce fichier ({ext}): {result.stderr.strip()[:200]}",
            )

        with open(pdf_path, "rb") as f:
            return f.read()


def _extract_pdf(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        try:
            reader = PdfReader(tmp.name)
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(400, f"PDF illisible: {exc}")
        parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts).strip()


def _extract_docx(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        try:
            document = docx_lib.Document(tmp.name)
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(400, f"Fichier Word illisible: {exc}")
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def _extract_odt(content: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".odt", delete=True) as tmp:
        tmp.write(content)
        tmp.flush()
        try:
            doc = load_odf(tmp.name)
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(400, f"Fichier ODT illisible: {exc}")
        paragraphs = doc.getElementsByType(OdfParagraph)
        parts = [teletype.extractText(p) for p in paragraphs]
    return "\n".join(parts).strip()


def _extract_rtf(content: bytes) -> str:
    try:
        raw = content.decode("utf-8", errors="ignore")
        return rtf_to_text(raw).strip()
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(400, f"Fichier RTF illisible: {exc}")


def _extract_plain_text(content: bytes) -> str:
    try:
        return content.decode("utf-8").strip()
    except UnicodeDecodeError:
        try:
            return content.decode("latin-1").strip()
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(400, f"Fichier texte illisible: {exc}")


def _extract_via_libreoffice(content: bytes, ext: str) -> str:
    """Formats anciens/binaires (.doc, .xls, .ppt) et variantes modernes
    (.xlsx, .pptx, .ods, .odp) : convertis en PDF via LibreOffice headless,
    puis extraits comme un PDF natif."""
    pdf_bytes = _convert_with_libreoffice(content, ext)
    return _extract_pdf(pdf_bytes)


EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".odt": _extract_odt,
    ".rtf": _extract_rtf,
    ".txt": _extract_plain_text,
    ".doc": _extract_via_libreoffice,
    ".xls": _extract_via_libreoffice,
    ".ppt": _extract_via_libreoffice,
    ".xlsx": _extract_via_libreoffice,
    ".pptx": _extract_via_libreoffice,
    ".ods": _extract_via_libreoffice,
    ".odp": _extract_via_libreoffice,
    ".md": _extract_plain_text,
}


def extract_text(content: bytes, filename: str) -> str:
    if len(content) > MAX_FILE_BYTES:
        raise HTTPException(400, "Fichier trop volumineux (20 Mo max)")

    ext = os.path.splitext(filename.lower())[1]
    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        raise HTTPException(400, f"Type de fichier non supporté: {ext}")

    if ext in LIBREOFFICE_EXTENSIONS:
        text = extractor(content, ext)
    else:
        text = extractor(content)

    if not text:
        raise HTTPException(
            400,
            "Impossible d'extraire du texte de ce fichier "
            "(document scanné/image, OCR non encore supporté)",
        )
    return text[:MAX_TEXT_CHARS]


def call_mistral(prompt: str, temperature: float = 0.3) -> str:
    if not MISTRAL_API_KEY:
        raise HTTPException(500, "Service mal configuré (clé Mistral manquante côté serveur)")
    try:
        resp = requests.post(
            "https://api.mistral.ai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {MISTRAL_API_KEY}",
                "Content-Type": "application/json",
            },
            json={
                "model": MISTRAL_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": temperature,
            },
            timeout=90,
        )
        resp.raise_for_status()
    except requests.RequestException as exc:
        raise HTTPException(502, f"Erreur du service IA: {exc}")

    data = resp.json()
    return data["choices"][0]["message"]["content"].strip()


def check_supported_upload(file: UploadFile) -> None:
    ext = os.path.splitext((file.filename or "").lower())[1]
    if not file.filename or ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise HTTPException(400, f"Type de fichier non supporté (formats acceptés: {supported})")


async def read_upload_limited(file: UploadFile, max_bytes: int = MAX_FILE_BYTES) -> bytes:
    """Lit le corps d'un UploadFile par blocs et coupe DES QUE la taille
    depasse max_bytes, sans attendre la fin du transfert. Remplace
    `await file.read()` (qui lit tout avant toute verification) : sans ce
    garde-fou, un fichier de plusieurs Go passe entierement par le reseau et
    le disque avant d'etre rejete par extract_text (deni de service facile).
    LimitRequestBody cote Apache existe aussi en defense complementaire, mais
    ne s'applique pas de facon fiable aux routes proxifiees (verifie en
    test : une requete au-dessus de la limite Apache configuree passe tout
    de meme jusqu'au backend via mod_proxy) — donc le vrai garde-fou doit
    etre ici, cote application."""
    chunk_size = 1024 * 1024
    total = 0
    chunks = []
    while True:
        chunk = await file.read(chunk_size)
        if not chunk:
            break
        total += len(chunk)
        if total > max_bytes:
            raise HTTPException(400, "Fichier trop volumineux (20 Mo max)")
        chunks.append(chunk)
    return b"".join(chunks)


@app.get("/health")
def health():
    return {"ok": True}


@app.get("/version")
def version():
    """Utilisé par le client Linux installé pour se comparer à la dernière version
    disponible et afficher une notification de mise à jour."""
    return {
        "version": CURRENT_VERSION,
        "download_url": f"https://docpilote.prestalibre.org/downloads/docpilote_{CURRENT_VERSION}_amd64.deb",
    }


@app.get("/version-windows")
def version_windows():
    """Equivalent de /version pour le client Windows (installeur .exe distinct
    du .deb Linux, numérotation de version indépendante)."""
    return {
        "version": CURRENT_VERSION_WINDOWS,
        "download_url": f"https://docpilote.prestalibre.org/downloads/DocPilote-Setup-{CURRENT_VERSION_WINDOWS}.exe",
    }


@app.get("/stats/public")
def stats_public():
    """Compteur public de téléchargements, affiché sur la landing page.
    Aucune donnée personnelle exposée."""
    with db_conn() as conn:
        total = conn.execute("SELECT COUNT(*) FROM downloads").fetchone()[0]
    return {"total_downloads": total}


@app.post("/billing/checkout")
def billing_checkout(email: str = Form("")):
    """Cree une session Stripe Checkout pour l'abonnement mensuel DocPilote Pro.
    Ne cree PAS la licence ici : la licence n'est emise qu'a la confirmation
    reelle du paiement via le webhook (evite les licences 'fantomes' si
    l'utilisateur ferme la page Stripe sans payer)."""
    if not STRIPE_SECRET_KEY or not STRIPE_PRICE_ID:
        raise HTTPException(503, "Paiement non configure pour le moment.")
    try:
        session = stripe.checkout.Session.create(
            mode="subscription",
            line_items=[{"price": STRIPE_PRICE_ID, "quantity": 1}],
            customer_email=email.strip() or None,
            success_url=CHECKOUT_SUCCESS_URL,
            cancel_url=CHECKOUT_CANCEL_URL,
            allow_promotion_codes=True,
        )
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(502, f"Erreur Stripe: {exc}") from exc
    return {"checkout_url": session.url}


@app.post("/billing/webhook")
async def billing_webhook(request: Request):
    """Recoit les evenements Stripe. Emet/active/desactive les licences selon
    le cycle de vie reel de l'abonnement (paiement confirme, echec, resiliation).
    Signature verifiee obligatoirement (pas de traitement d'evenement non signe)."""
    if not STRIPE_WEBHOOK_SECRET:
        raise HTTPException(503, "Webhook non configure.")
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature", "")
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except (ValueError, stripe.error.SignatureVerificationError) as exc:
        raise HTTPException(400, f"Webhook invalide: {exc}") from exc

    etype = event["type"]
    data = event["data"]["object"]

    if etype == "checkout.session.completed":
        customer_id = data.get("customer")
        subscription_id = data.get("subscription")
        email = (data.get("customer_details") or {}).get("email") or data.get("customer_email") or ""
        license_key = generate_license_key()
        period_end = datetime.now(timezone.utc) + timedelta(days=35)
        with db_conn() as conn:
            conn.execute(
                """
                INSERT INTO licenses (license_key, email, status, stripe_customer_id, stripe_subscription_id, current_period_end, created_at)
                VALUES (?, ?, 'active', ?, ?, ?, ?)
                """,
                (license_key, email, customer_id, subscription_id, period_end.isoformat(), datetime.now(timezone.utc).isoformat()),
            )
        # TODO email d'envoi de la cle non branche (pas de service SMTP
        # transactionnel configure pour ce projet) -> la cle est affichee
        # directement sur la page de succes cote client (session_id Stripe).

    elif etype == "invoice.payment_succeeded":
        subscription_id = data.get("subscription")
        if subscription_id:
            period_end = datetime.now(timezone.utc) + timedelta(days=35)
            with db_conn() as conn:
                conn.execute(
                    "UPDATE licenses SET status = 'active', current_period_end = ? WHERE stripe_subscription_id = ?",
                    (period_end.isoformat(), subscription_id),
                )

    elif etype in ("invoice.payment_failed", "customer.subscription.deleted"):
        subscription_id = data.get("subscription") or data.get("id")
        if subscription_id:
            with db_conn() as conn:
                conn.execute(
                    "UPDATE licenses SET status = 'suspended' WHERE stripe_subscription_id = ?",
                    (subscription_id,),
                )

    return {"received": True}


@app.get("/billing/session/{session_id}")
def billing_session_lookup(session_id: str):
    """Retrouve la cle de licence juste emise a partir du session_id Stripe,
    pour l'afficher sur la page de succes cote client (aucune donnee bancaire
    ne transite jamais par ce backend, uniquement gere par Stripe)."""
    if not STRIPE_SECRET_KEY:
        raise HTTPException(503, "Paiement non configure.")
    try:
        session = stripe.checkout.Session.retrieve(session_id)
    except Exception as exc:  # noqa: BLE001
        raise HTTPException(404, f"Session introuvable: {exc}") from exc
    customer_id = session.get("customer")
    if not customer_id:
        raise HTTPException(404, "Session sans client associe.")
    with db_conn() as conn:
        row = conn.execute(
            "SELECT license_key FROM licenses WHERE stripe_customer_id = ? ORDER BY id DESC LIMIT 1",
            (customer_id,),
        ).fetchone()
    if not row:
        raise HTTPException(404, "Licence pas encore emise (le paiement peut prendre quelques secondes a se confirmer).")
    return {"license_key": row[0]}


@app.get("/license/status")
def license_status(key: str = "", device_id: str = ""):
    """Verifie la validite d'une licence et, si device_id est fourni, tente
    la liaison machine a ce moment-la (plutot que d'attendre le premier
    appel /process etc.) pour un message d'erreur clair a l'activation."""
    valid = license_is_valid(key)
    if not valid:
        return {"valid": False}
    try:
        check_device_binding(key, device_id)
    except HTTPException as exc:
        return {"valid": False, "device_conflict": True, "message": str(exc.detail)}
    return {"valid": True}


@app.get("/downloads/{filename}")
async def tracked_download(filename: str, request: Request):
    """Sert les .deb (Linux) et .exe (Windows) en journalisant le
    téléchargement (RGPD: IP hashée, jamais stockée en clair). Bloque toute
    tentative de path traversal."""
    safe_name = os.path.basename(filename)
    is_linux_deb = safe_name.startswith("docpilote_") and safe_name.endswith("_amd64.deb")
    is_windows_exe = safe_name.startswith("DocPilote-Setup-") and safe_name.endswith(".exe")
    if safe_name != filename or not (is_linux_deb or is_windows_exe):
        raise HTTPException(404, "Fichier introuvable")

    path = os.path.join(DOWNLOADS_DIR, safe_name)
    if not os.path.isfile(path):
        raise HTTPException(404, "Fichier introuvable")

    if is_linux_deb:
        version_part = safe_name.replace("docpilote_", "").replace("_amd64.deb", "")
        media_type = "application/vnd.debian.binary-package"
    else:
        version_part = "win-" + safe_name.replace("DocPilote-Setup-", "").replace(".exe", "")
        media_type = "application/vnd.microsoft.portable-executable"
    log_download(
        version_part,
        _client_ip(request),
        request.headers.get("user-agent", ""),
        request.headers.get("cf-ipcountry", ""),
        request.headers.get("cf-ipcity", ""),
        request.headers.get("cf-region", ""),
        request.headers.get("cf-iplatitude", ""),
        request.headers.get("cf-iplongitude", ""),
    )

    return FileResponse(path, media_type=media_type, filename=safe_name)


@app.post("/process")
async def process(request: Request, file: UploadFile = File(...), plugin: str = Form("summarize"), license_key: str = Form(""), device_id: str = Form(""), lang: str = Form("fr")):
    started = time.monotonic()
    ext = os.path.splitext((file.filename or "").lower())[1]
    try:
        if plugin not in PLUGIN_PROMPTS:
            raise HTTPException(400, f"Plugin inconnu: {plugin}")
        check_quota(request, license_key, device_id)
        check_supported_upload(file)

        content = await read_upload_limited(file)
        text = extract_text(content, file.filename)

        prompt = PLUGIN_PROMPTS[plugin].format(label=lang_label(lang), text=text)
        result = call_mistral(prompt)
    except HTTPException as exc:
        log_usage("process", plugin, ext, False, str(exc.detail), int((time.monotonic() - started) * 1000), _client_ip(request))
        raise
    log_usage("process", plugin, ext, True, "", int((time.monotonic() - started) * 1000), _client_ip(request))
    return {"result": result}


@app.post("/extract")
async def extract(request: Request, file: UploadFile = File(...), license_key: str = Form(""), device_id: str = Form(""), lang: str = Form("fr")):
    started = time.monotonic()
    ext = os.path.splitext((file.filename or "").lower())[1]
    try:
        check_quota(request, license_key, device_id)
        check_supported_upload(file)
        content = await read_upload_limited(file)
        text = extract_text(content, file.filename)

        prompt = EXTRACT_PROMPT.format(label=lang_label(lang), text=text)
        result = call_mistral(prompt, temperature=0.1)

        # Nettoyage basique si le modèle a entouré le JSON de balises markdown
        cleaned = result.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[-1]
            if cleaned.endswith("```"):
                cleaned = cleaned.rsplit("```", 1)[0]
    except HTTPException as exc:
        log_usage("extract", None, ext, False, str(exc.detail), int((time.monotonic() - started) * 1000), _client_ip(request))
        raise
    log_usage("extract", None, ext, True, "", int((time.monotonic() - started) * 1000), _client_ip(request))
    return {"result": cleaned.strip()}


@app.post("/compare")
async def compare(request: Request, file1: UploadFile = File(...), file2: UploadFile = File(...), license_key: str = Form(""), device_id: str = Form(""), lang: str = Form("fr")):
    started = time.monotonic()
    ext = os.path.splitext((file1.filename or "").lower())[1]
    try:
        check_quota(request, license_key, device_id)
        check_supported_upload(file1)
        check_supported_upload(file2)

        content1 = await read_upload_limited(file1)
        content2 = await read_upload_limited(file2)
        text1 = extract_text(content1, file1.filename)
        text2 = extract_text(content2, file2.filename)

        prompt = COMPARE_PROMPT.format(label=lang_label(lang), text1=text1, text2=text2)
        result = call_mistral(prompt)
    except HTTPException as exc:
        log_usage("compare", None, ext, False, str(exc.detail), int((time.monotonic() - started) * 1000), _client_ip(request))
        raise
    log_usage("compare", None, ext, True, "", int((time.monotonic() - started) * 1000), _client_ip(request))
    return {"result": result}


@app.post("/ask")
async def ask(request: Request, file: UploadFile = File(...), question: str = Form(...), license_key: str = Form(""), device_id: str = Form(""), lang: str = Form("fr")):
    started = time.monotonic()
    ext = os.path.splitext((file.filename or "").lower())[1]
    try:
        check_quota(request, license_key, device_id)
        check_supported_upload(file)
        if not question.strip():
            raise HTTPException(400, "Question vide")

        content = await read_upload_limited(file)
        text = extract_text(content, file.filename)

        prompt = ASK_PROMPT.format(label=lang_label(lang), question=question.strip(), text=text)
        result = call_mistral(prompt)
    except HTTPException as exc:
        log_usage("ask", None, ext, False, str(exc.detail), int((time.monotonic() - started) * 1000), _client_ip(request))
        raise
    log_usage("ask", None, ext, True, "", int((time.monotonic() - started) * 1000), _client_ip(request))
    return {"result": result}


@app.post("/transform")
async def transform(request: Request, file: UploadFile = File(...), target: str = Form(...), license_key: str = Form(""), device_id: str = Form(""), lang: str = Form("fr")):
    started = time.monotonic()
    ext = os.path.splitext((file.filename or "").lower())[1]
    try:
        check_quota(request, license_key, device_id)
        check_supported_upload(file)
        if target not in TRANSFORM_TARGETS:
            raise HTTPException(400, f"Format de transformation inconnu: {target}")

        content = await read_upload_limited(file)
        text = extract_text(content, file.filename)

        prompt = (
            f"Transforme ce document en {TRANSFORM_TARGETS[target]}. "
            f"Réponds en {lang_label(lang)}, directement avec le résultat, sans commentaire méta "
            f"autour.\n\nDocument:\n{text}"
        )
        result = call_mistral(prompt)
    except HTTPException as exc:
        log_usage("transform", target, ext, False, str(exc.detail), int((time.monotonic() - started) * 1000), _client_ip(request))
        raise
    log_usage("transform", target, ext, True, "", int((time.monotonic() - started) * 1000), _client_ip(request))
    return {"result": result}


# ---------------------------------------------------------------------------
# Panel admin (Basic Auth) : téléchargements, usage, erreurs
# ---------------------------------------------------------------------------


@app.get("/admin/stats")
def admin_stats(_: bool = Depends(require_admin)):
    with db_conn() as conn:
        total_downloads = conn.execute("SELECT COUNT(*) FROM downloads").fetchone()[0]
        downloads_by_version = conn.execute(
            "SELECT version, COUNT(*) FROM downloads GROUP BY version ORDER BY version DESC"
        ).fetchall()
        downloads_by_day = conn.execute(
            """
            SELECT substr(created_at, 1, 10) as day, COUNT(*) FROM downloads
            GROUP BY day ORDER BY day DESC LIMIT 30
            """
        ).fetchall()
        unique_downloaders = conn.execute(
            "SELECT COUNT(DISTINCT ip_hash) FROM downloads"
        ).fetchone()[0]
        downloads_by_country = conn.execute(
            """
            SELECT COALESCE(country, '??'), COUNT(*) FROM downloads
            GROUP BY country ORDER BY COUNT(*) DESC
            """
        ).fetchall()
        downloads_by_city = conn.execute(
            """
            SELECT city, region, country, AVG(latitude), AVG(longitude), COUNT(*)
            FROM downloads
            WHERE city IS NOT NULL AND latitude IS NOT NULL AND longitude IS NOT NULL
            GROUP BY city, country ORDER BY COUNT(*) DESC
            """
        ).fetchall()

        total_usage = conn.execute("SELECT COUNT(*) FROM usage_events").fetchone()[0]
        usage_success = conn.execute(
            "SELECT COUNT(*) FROM usage_events WHERE success = 1"
        ).fetchone()[0]
        usage_by_endpoint = conn.execute(
            """
            SELECT endpoint, plugin, COUNT(*), SUM(success), AVG(duration_ms)
            FROM usage_events GROUP BY endpoint, plugin ORDER BY COUNT(*) DESC
            """
        ).fetchall()
        usage_by_ext = conn.execute(
            "SELECT ext, COUNT(*) FROM usage_events GROUP BY ext ORDER BY COUNT(*) DESC"
        ).fetchall()
        usage_by_day = conn.execute(
            """
            SELECT substr(created_at, 1, 10) as day, COUNT(*), SUM(success) FROM usage_events
            GROUP BY day ORDER BY day DESC LIMIT 30
            """
        ).fetchall()
        unique_users = conn.execute(
            "SELECT COUNT(DISTINCT ip_hash) FROM usage_events"
        ).fetchone()[0]
        recent_errors = conn.execute(
            """
            SELECT endpoint, plugin, ext, error_msg, created_at FROM usage_events
            WHERE success = 0 ORDER BY created_at DESC LIMIT 30
            """
        ).fetchall()

        total_licenses = conn.execute("SELECT COUNT(*) FROM licenses").fetchone()[0]
        active_licenses = conn.execute(
            "SELECT COUNT(*) FROM licenses WHERE status = 'active'"
        ).fetchone()[0]
        recent_licenses = conn.execute(
            "SELECT license_key, email, status, current_period_end, created_at FROM licenses ORDER BY created_at DESC LIMIT 30"
        ).fetchall()

    return {
        "downloads": {
            "total": total_downloads,
            "unique_ip_hashes": unique_downloaders,
            "by_version": [{"version": v, "count": c} for v, c in downloads_by_version],
            "by_day": [{"day": d, "count": c} for d, c in downloads_by_day],
            "by_country": [{"country": c, "count": n} for c, n in downloads_by_country],
            "by_city": [
                {"city": city, "region": region, "country": country, "lat": lat, "lon": lon, "count": n}
                for city, region, country, lat, lon, n in downloads_by_city
            ],
        },
        "usage": {
            "total": total_usage,
            "success": usage_success,
            "failure": total_usage - usage_success,
            "unique_ip_hashes": unique_users,
            "by_endpoint": [
                {"endpoint": e, "plugin": p, "count": c, "success": s or 0, "avg_ms": round(a or 0)}
                for e, p, c, s, a in usage_by_endpoint
            ],
            "by_ext": [{"ext": e, "count": c} for e, c in usage_by_ext],
            "by_day": [{"day": d, "count": c, "success": s or 0} for d, c, s in usage_by_day],
            "recent_errors": [
                {"endpoint": e, "plugin": p, "ext": ex, "error": err, "at": at}
                for e, p, ex, err, at in recent_errors
            ],
        },
        "current_version": CURRENT_VERSION,
        "licenses": {
            "total": total_licenses,
            "active": active_licenses,
            "mrr_estimate_eur": active_licenses * MONTHLY_PRICE_EUR,
            "recent": [
                {"key": k, "email": e, "status": s, "period_end": pe, "created_at": ca}
                for k, e, s, pe, ca in recent_licenses
            ],
        },
    }


@app.get("/admin", response_class=HTMLResponse)
def admin_panel(_: bool = Depends(require_admin)):
    return ADMIN_HTML


ADMIN_HTML = """<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>DocPilote — Admin</title>
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/jsvectormap@1.7.0/dist/jsvectormap.min.css">
<link rel="stylesheet" href="https://unpkg.com/leaflet@1.9.4/dist/leaflet.css">
<style>
  :root{--bg:#0b0f14;--card:#141c27;--border:#223042;--text:#e8eef4;--muted:#9fb0c3;--accent:#4fd1c5;--accent2:#7c9cff;--bad:#ff8e8e;}
  *{box-sizing:border-box;}
  body{margin:0;font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;background:var(--bg);color:var(--text);padding:24px;}
  h1{font-size:1.4rem;margin:0 0 4px;}
  .sub{color:var(--muted);margin:0 0 24px;font-size:.9rem;}
  .grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:16px;margin-bottom:24px;}
  .card{background:var(--card);border:1px solid var(--border);border-radius:10px;padding:18px;}
  .card .big{font-size:1.8rem;font-weight:800;background:linear-gradient(135deg,var(--accent),var(--accent2));-webkit-background-clip:text;background-clip:text;color:transparent;}
  .card .label{color:var(--muted);font-size:.85rem;margin-top:4px;}
  table{width:100%;border-collapse:collapse;font-size:.85rem;}
  th,td{text-align:left;padding:8px 10px;border-bottom:1px solid var(--border);}
  th{color:var(--muted);font-weight:600;}
  section{margin-bottom:32px;}
  section h2{font-size:1.05rem;margin:0 0 12px;color:var(--text);}
  .err{color:var(--bad);}
  .ok{color:var(--accent);}
  .refresh{color:var(--muted);font-size:.8rem;}
  #map{width:100%;height:420px;background:transparent;}
  #city-map{width:100%;height:420px;background:transparent;border-radius:8px;}
  .leaflet-popup-content-wrapper{background:#141c27;color:#e8edf2;}
  .leaflet-popup-tip{background:#141c27;}
  .jvm-zoom-btn{background:var(--card);border:1px solid var(--border);color:var(--text);}
  .jvm-tooltip{background:var(--card);border:1px solid var(--border);color:var(--text);}
</style>
</head>
<body>
<h1>DocPilote — Tableau de bord</h1>
<p class="sub">Version distribuée actuelle : <strong id="cv">...</strong> &nbsp;·&nbsp; <span class="refresh">Actualisé automatiquement toutes les 30s</span></p>

<div class="grid" id="kpis"></div>

<section>
<h2>Téléchargements par pays</h2>
<div id="map"></div>
</section>

<section>
<h2>Villes des téléchargeurs</h2>
<p class="sub" style="margin:0 0 12px;">Géolocalisation approximative basée sur l'IP (Cloudflare), jamais l'IP elle-même. Taille du point proportionnelle au nombre de téléchargements.</p>
<div id="city-map"></div>
</section>

<section>
<h2>Téléchargements par version</h2>
<table id="dl-version"><thead><tr><th>Version</th><th>Téléchargements</th></tr></thead><tbody></tbody></table>
</section>

<section>
<h2>Téléchargements (30 derniers jours)</h2>
<table id="dl-day"><thead><tr><th>Jour</th><th>Téléchargements</th></tr></thead><tbody></tbody></table>
</section>

<section>
<h2>Usage par action</h2>
<table id="usage-endpoint"><thead><tr><th>Endpoint</th><th>Plugin</th><th>Appels</th><th>Succès</th><th>Durée moy.</th></tr></thead><tbody></tbody></table>
</section>

<section>
<h2>Usage par format de fichier</h2>
<table id="usage-ext"><thead><tr><th>Format</th><th>Appels</th></tr></thead><tbody></tbody></table>
</section>

<section>
<h2>Licences (abonnements Stripe)</h2>
<table id="licenses"><thead><tr><th>Cle</th><th>Email</th><th>Statut</th><th>Fin de periode</th><th>Depuis</th></tr></thead><tbody></tbody></table>
</section>

<section>
<h2>Erreurs récentes</h2>
<table id="errors"><thead><tr><th>Quand</th><th>Endpoint</th><th>Plugin</th><th>Format</th><th>Erreur</th></tr></thead><tbody></tbody></table>
</section>

<script src="https://cdn.jsdelivr.net/npm/jsvectormap@1.7.0/dist/jsvectormap.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/jsvectormap@1.7.0/dist/maps/world.js"></script>
<script src="https://unpkg.com/leaflet@1.9.4/dist/leaflet.js"></script>
<script>
let worldMap = null;
let cityMap = null;
let cityMarkers = [];

function renderCityMap(byCity){
  if(!cityMap){
    cityMap = L.map('city-map', {scrollWheelZoom:false}).setView([46.6, 2.5], 5);
    L.tileLayer('https://{s}.basemaps.cartocdn.com/dark_all/{z}/{x}/{y}{r}.png', {
      attribution: '&copy; OpenStreetMap &copy; CARTO',
      maxZoom: 18,
    }).addTo(cityMap);
  }
  for(const m of cityMarkers){ cityMap.removeLayer(m); }
  cityMarkers = [];
  if(!byCity || byCity.length === 0) return;
  const maxCount = Math.max(...byCity.map(r => r.count));
  for(const row of byCity){
    if(row.lat == null || row.lon == null) continue;
    const radius = 6 + (row.count / maxCount) * 22;
    const marker = L.circleMarker([row.lat, row.lon], {
      radius,
      color: '#4fd1c5',
      weight: 1,
      fillColor: '#4fd1c5',
      fillOpacity: 0.45,
    }).addTo(cityMap);
    marker.bindPopup(`<strong>${row.city}</strong>${row.region ? ' — ' + row.region : ''}<br>${row.count} téléchargement${row.count > 1 ? 's' : ''}`);
    cityMarkers.push(marker);
  }
}

function renderMap(byCountry){
  const values = {};
  for(const row of byCountry){
    if(row.country && row.country !== '??'){
      values[row.country] = row.count;
    }
  }
  const hasData = Object.keys(values).length > 0;
  if(!worldMap){
    worldMap = new jsVectorMap({
      selector: '#map',
      map: 'world',
      backgroundColor: 'transparent',
      regionStyle: { initial: { fill: '#223042' } },
      series: hasData ? { regions: [{
        values,
        scale: ['#1a3a35', '#4fd1c5'],
        normalizeFunction: 'polynomial'
      }] } : {},
    });
  } else if(hasData) {
    worldMap.series.regions[0].setValues(values);
  }
}

async function load(){
  const res = await fetch('/api/admin/stats', {credentials:'include'});
  if(!res.ok){ document.body.innerHTML = '<p style="color:#ff8e8e">Erreur de chargement des stats.</p>'; return; }
  const d = await res.json();

  document.getElementById('cv').textContent = 'v' + d.current_version;

  const kpis = document.getElementById('kpis');
  kpis.innerHTML = '';
  const items = [
    [d.downloads.total, 'Téléchargements totaux'],
    [d.downloads.unique_ip_hashes, 'Téléchargeurs uniques (approx.)'],
    [d.usage.total, 'Appels API totaux'],
    [d.usage.success + ' / ' + d.usage.total, 'Appels réussis'],
    [d.usage.unique_ip_hashes, 'Utilisateurs actifs uniques (approx.)'],
    [d.licenses.active + ' / ' + d.licenses.total, 'Licences actives'],
    [d.licenses.mrr_estimate_eur + ' €', 'MRR estimé'],
  ];
  for(const [big, label] of items){
    const c = document.createElement('div');
    c.className = 'card';
    c.innerHTML = '<div class="big">'+big+'</div><div class="label">'+label+'</div>';
    kpis.appendChild(c);
  }

  const dlVersionBody = document.querySelector('#dl-version tbody');
  renderMap(d.downloads.by_country || []);
  renderCityMap(d.downloads.by_city || []);

  dlVersionBody.innerHTML = d.downloads.by_version.map(r => `<tr><td>v${r.version}</td><td>${r.count}</td></tr>`).join('') || '<tr><td colspan=2>Aucune donnée</td></tr>';

  const dlDayBody = document.querySelector('#dl-day tbody');
  dlDayBody.innerHTML = d.downloads.by_day.map(r => `<tr><td>${r.day}</td><td>${r.count}</td></tr>`).join('') || '<tr><td colspan=2>Aucune donnée</td></tr>';

  const usageEpBody = document.querySelector('#usage-endpoint tbody');
  usageEpBody.innerHTML = d.usage.by_endpoint.map(r => `<tr><td>${r.endpoint}</td><td>${r.plugin || '—'}</td><td>${r.count}</td><td class="${r.success===r.count ? 'ok':''}">${r.success}/${r.count}</td><td>${r.avg_ms} ms</td></tr>`).join('') || '<tr><td colspan=5>Aucune donnée</td></tr>';

  const usageExtBody = document.querySelector('#usage-ext tbody');
  usageExtBody.innerHTML = d.usage.by_ext.map(r => `<tr><td>${r.ext || '—'}</td><td>${r.count}</td></tr>`).join('') || '<tr><td colspan=2>Aucune donnée</td></tr>';

  const licBody = document.querySelector('#licenses tbody');
  licBody.innerHTML = d.licenses.recent.map(r => `<tr><td>${r.key}</td><td>${r.email || '—'}</td><td class="${r.status==='active'?'ok':'err'}">${r.status}</td><td>${(r.period_end||'').replace('T',' ').slice(0,19) || '—'}</td><td>${r.created_at.replace('T',' ').slice(0,19)}</td></tr>`).join('') || '<tr><td colspan=5>Aucune licence pour le moment</td></tr>';

  const errBody = document.querySelector('#errors tbody');
  errBody.innerHTML = d.usage.recent_errors.map(r => `<tr><td>${r.at.replace('T',' ').slice(0,19)}</td><td>${r.endpoint}</td><td>${r.plugin || '—'}</td><td>${r.ext || '—'}</td><td class="err">${r.error}</td></tr>`).join('') || '<tr><td colspan=5>Aucune erreur 🎉</td></tr>';
}
load();
setInterval(load, 30000);
</script>
</body>
</html>
"""

